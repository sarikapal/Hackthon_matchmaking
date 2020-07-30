# -*- coding: utf-8 -*-
"""
Created on Wed Jul 22 16:58:56 2020

@author: sarpa
"""

import gensim
#import PyPDF2
#import slate3k as slate
import os
#from pathlib import Path
#import nltk
import re
import glob
 
#from gensim.models import Doc2Vec
from gensim.models.doc2vec import TaggedDocument
from docx import Document
from nltk import word_tokenize
from nltk.stem import PorterStemmer
from nltk.corpus import stopwords
#import smart_open

import requests
import time

#****Clean and preprocess the text
def clean_text(text):
    '''
    Removes default bad characters
    '''
#     if not (pd.isnull(text)):
    # text = filter(lambda x: x in string.printable, text)
    bad_chars = set(["\n", "@", "+", "'", '"', '\\','(',')', '', '\\n', '', '?', '#', ',', '[',']', '%', '$', '&', ';', '!', ';', ':',"*", "_", "=", "}", "{"])
    for char in bad_chars:
        text = text.replace(char, "")
    text = re.sub('\d+', "", text)
    text = re.sub(" +", " ", text)
    text = text.replace(".", " ")
    text = text.replace("/", " ")
    return text
 
    
def stop_and_stem(text, stem=False, stemmer = PorterStemmer()):
    '''
    Removes stopwords and does stemming
    '''
    stoplist = stopwords.words('english')
    if stem:
        text_stemmed = [stemmer.stem(word) for word in word_tokenize(text) if word not in stoplist and len(word) > 3]
    else:
        text_stemmed = [word for word in word_tokenize(text) if word not in stoplist and len(word) > 3]
    text = ' '.join(text_stemmed)
    return text
    

#********Import resume in a file set************************************
path="C:/Users/sarpa/OneDrive - Microsoft/Wipro/Text Analytics/Resume match"
fileset = [file for file in glob.glob(path + "/Resume/**/*.pdf", recursive=True)]

#******use microsoft read api to extract the text from documents

url="https://eastus.api.cognitive.microsoft.com/"
key=""



def callReadapi(file,url,key):
    with open(file, "rb") as image_stream:
        
        data=image_stream.read()
        vision_base_url=url
        text_recognition_url = vision_base_url + "vision/v3.0-preview/read/analyze?language=en"
        headers  = {'Ocp-Apim-Subscription-Key': key,
                   "Content-Type": "application/octet-stream"}
        #params   = {'handwriting' : True}
        params   = {'mode' : "Printed"}
        response = requests.post(text_recognition_url, headers=headers, params=params, data=data)
        response.raise_for_status()
        analysis1 = {}
        while not "analyzeResult" in analysis1:
            response_final = requests.get(response.headers["Operation-Location"], headers=headers)
            analysis1       = response_final.json()
            time.sleep(2)
        
        result_data = analysis1["analyzeResult"]["readResults"]    
        return result_data
    
    
import itertools
corpus = []
rawcorpus=[]
filenameraw=[]
summarizedtext=[]
tag_remove = fileset
#model = Summarizer()
for doc in fileset:
    print(doc)
    try:
        if doc.split(".")[1] == "pdf":
            try:
                extractedtext=callReadapi(doc,url,key) 
            except:
                print("Error calling rest api")
            try:
                texts=[]
                for i in range(len(extractedtext)):
                    text=[ii["text"] for ii in extractedtext[i]["lines"]]
                    texts.append('.'.join(text))
            
                cont_text=' '. join(list(itertools.chain(texts)))
                cont_text= clean_text(cont_text)
                    #summary=model(cont_text)
                corpus.append(stop_and_stem(cont_text))
                    #summarizedtext.append(summary) 
            except:
                print("error in extraction")
           
            filenameraw.append([os.path.split(doc)[1].split('.')[0]])
            #print(len(corpus))
        elif doc.split(".")[1] == "docx":
            doc = Document(doc)
            text = ''
            for para in doc.paragraphs:
                text += para.text
            #summary=model(text)
            corpus.append(text) #### stop and stem not applied here
            #summarizedtext.append(summary)        
            filenameraw.append([os.path.split(doc)[1].split('.')[0]])

    except:
        print(doc)





#*********************************************************************************

 


 


 

#****************creating tagged corpus
tagged_cr = []
for idx, doc in enumerate(corpus):
    tagged_cr.append(TaggedDocument(words=doc.split(), tags=[idx]))
    
    
#*********model prep    
model = gensim.models.doc2vec.Doc2Vec(vector_size=100, min_count=1, epochs=80, alpha=0.025)
model.build_vocab(tagged_cr)
model.train(tagged_cr, total_examples=model.corpus_count, epochs=model.epochs)

 
#vector = model.infer_vector(['dancing', 'reading', 'theatre', 'machine', 'learning', 'skills'])
#print(vector)

##Overfitting Test
ranks = []
second_ranks = []
for doc_id in range(len(tagged_cr)):
    inferred_vector = model.infer_vector(tagged_cr[doc_id].words)
    sims = model.docvecs.most_similar([inferred_vector], topn=len(model.docvecs))
#     print(sims)
#    rank = sims[0][0]
#    ranks.append(rank)
    rank = [docid for docid, sim in sims].index(doc_id)
    ranks.append(rank)
    second_ranks.append(sims[1])
print("Top Matching document", ranks)
print("Second Rank\n", second_ranks)

 


   
#**********************with API******************
test_doc = []
text=""
fp=(path + "/JD/JD_CA1.pdf")
extracted_testtext=callReadapi(fp,url,key)

texts=[]
for i in range(len(extracted_testtext)):
    text=[ii["text"] for ii in extracted_testtext[i]["lines"]]
    texts.append('.'.join(text))
            
    cont_text=' '. join(list(itertools.chain(texts)))
    cont_text= clean_text(cont_text)
    
text = clean_text(cont_text)
test_doc.append(stop_and_stem(text))

 

inferred_vector = model.infer_vector(test_doc[0].split())
model.docvecs.most_similar([inferred_vector], topn=len(model.docvecs))
